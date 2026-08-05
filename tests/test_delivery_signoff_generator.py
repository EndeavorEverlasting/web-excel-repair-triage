from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT

from triage.delivery_signoff import SignoffValidationError, generate_signoff
from triage.delivery_signoff import generator as generator_module
from triage.delivery_signoff.schema import MINIMUM_HEADING_POINTS

ROOT = Path(__file__).parents[1]
FIXTURES = Path(__file__).parent / "fixtures" / "delivery_signoff"


@pytest.mark.parametrize(
    ("fixture", "expected_rows", "expected_poc"),
    [
        ("melville_3hq_20260805.json", 16, "Timothy W. Seeberger"),
        ("huntington_hospital_20260805.json", 3, "Morgan Mitchell"),
    ],
)
def test_generates_valid_equipment_only_signoff(tmp_path: Path, fixture: str, expected_rows: int, expected_poc: str) -> None:
    result = generate_signoff(FIXTURES / fixture, tmp_path / "delivery-signoff")
    manifest = _manifest(result.manifest_path)
    assert result.docx_path.is_file()
    assert result.preview_pdf_path.is_file()
    assert result.validation_log_path.read_text(encoding="utf-8").startswith("PASS")
    assert 1 <= manifest["page_count"] <= 2
    assert len(manifest["equipment_rows"]) == expected_rows
    assert manifest["serial_counts"] == {}
    assert manifest["serialized_assets_expected"] is False
    assert manifest["preview"]["page_hashes"]
    assert manifest["minimum_heading_points"] == MINIMUM_HEADING_POINTS
    assert expected_poc in _docx_xml(result.docx_path)


def test_small_serial_group_uses_portrait_and_exact_reconciliation(tmp_path: Path) -> None:
    payload = _serial_payload(["1", "11", "111"])
    source = _write_json(tmp_path / "serial-signoff.json", payload)
    result = generate_signoff(source, tmp_path / "out")
    manifest = _manifest(result.manifest_path)
    document = Document(result.docx_path)
    assert document.sections[0].orientation == WD_ORIENT.PORTRAIT
    assert manifest["layout_plan"]["serial_columns"] == 2
    assert manifest["serial_counts"]["Neuron"] == {"declared": 3, "rendered": 3, "duplicates": 0}


def test_long_or_dense_serial_group_uses_landscape(tmp_path: Path) -> None:
    serials = [f"LONG-SERIAL-NUMBER-{index:02d}-ABCDEFGHIJKLMN" for index in range(1, 7)]
    source = _write_json(tmp_path / "dense.json", _serial_payload(serials))
    result = generate_signoff(source, tmp_path / "out")
    manifest = _manifest(result.manifest_path)
    document = Document(result.docx_path)
    assert document.sections[0].orientation == WD_ORIENT.LANDSCAPE
    assert manifest["layout_plan"]["serial_columns"] == 2


def test_heading_runs_meet_layout_minimum(tmp_path: Path) -> None:
    result = generate_signoff(FIXTURES / "huntington_hospital_20260805.json", tmp_path / "out")
    document = Document(result.docx_path)
    required = {
        "Equipment / Stock Receipt",
        "Exceptions / Field Notes",
        "Receipt Acceptance",
    }
    found: dict[str, float] = {}
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in required and paragraph.runs:
            found[text] = min(run.font.size.pt for run in paragraph.runs if run.font.size is not None)
    assert set(found) == required
    assert all(size >= MINIMUM_HEADING_POINTS for size in found.values())


def test_rejects_invalid_quantity(tmp_path: Path) -> None:
    payload = _fixture_payload("huntington_hospital_20260805.json")
    payload["equipment_rows"][0]["quantity"] = -1
    source = _write_json(tmp_path / "bad.json", payload)
    with pytest.raises(SignoffValidationError, match="positive integer"):
        generate_signoff(source, tmp_path / "out")


def test_rejects_duplicate_serials(tmp_path: Path) -> None:
    source = _write_json(tmp_path / "bad-serials.json", _serial_payload(["1100000001", "1100000001"]))
    with pytest.raises(SignoffValidationError, match="duplicate serial"):
        generate_signoff(source, tmp_path / "out")


def test_rejects_duplicate_serialized_asset_types(tmp_path: Path) -> None:
    payload = _serial_payload(["1001"])
    payload["serialized_assets"].append(
        {"asset_type": "neuron", "identifiers": [{"serial_number": "1002"}]}
    )
    source = _write_json(tmp_path / "duplicate-group.json", payload)
    with pytest.raises(SignoffValidationError, match="duplicate serialized asset type"):
        generate_signoff(source, tmp_path / "out")


def test_rejects_serial_quantity_mismatch(tmp_path: Path) -> None:
    payload = _serial_payload(["1001", "1002"])
    payload["equipment_rows"][0]["quantity"] = 3
    source = _write_json(tmp_path / "quantity-mismatch.json", payload)
    with pytest.raises(SignoffValidationError, match="quantity mismatch"):
        generate_signoff(source, tmp_path / "out")


def test_rejects_incomplete_cable_identity(tmp_path: Path) -> None:
    payload = _fixture_payload("huntington_hospital_20260805.json")
    payload["equipment_rows"] = [
        {"equipment_type": "Ethernet Cable", "model_or_part": "CAT6", "color_or_variant": "", "quantity": 2}
    ]
    source = _write_json(tmp_path / "bad-cable.json", payload)
    with pytest.raises(SignoffValidationError, match="require model_or_part and color_or_variant"):
        generate_signoff(source, tmp_path / "out")


def test_failed_regeneration_preserves_last_valid_package(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    output_root = tmp_path / "out"
    source = FIXTURES / "huntington_hospital_20260805.json"
    first = generate_signoff(source, output_root)
    original_hash = first.docx_path.read_bytes()

    def fail_render(*_args: object, **_kwargs: object) -> object:
        raise SignoffValidationError("forced render failure")

    monkeypatch.setattr(generator_module, "render_docx", fail_render)
    with pytest.raises(SignoffValidationError, match="forced render failure"):
        generate_signoff(source, output_root)
    assert first.docx_path.read_bytes() == original_hash
    assert not list(first.package_dir.parent.glob(".*.tmp-*"))


def test_canonical_descendant_run_uses_outputs_backup_root() -> None:
    run_root = ROOT / "Outputs" / "delivery-signoff" / "review"
    assert generator_module._backup_root(run_root) == ROOT / "Outputs" / "backups" / "delivery-signoff"


def test_successful_regeneration_creates_backup(tmp_path: Path) -> None:
    output_root = tmp_path / "delivery-signoff"
    source = FIXTURES / "huntington_hospital_20260805.json"
    first = generate_signoff(source, output_root)
    first_hash = first.docx_path.read_bytes()
    second = generate_signoff(source, output_root)
    backups = list((tmp_path / "backups" / "delivery-signoff").rglob("input-spec.json"))
    assert backups
    assert second.docx_path.is_file()
    backed_up_docx = next(backups[0].parent.glob("*.docx"))
    assert backed_up_docx.read_bytes() == first_hash


def test_slug_collision_preserves_existing_package(tmp_path: Path) -> None:
    first_payload = _fixture_payload("huntington_hospital_20260805.json")
    first_payload["site"]["code"] = "A/B"
    first_payload["signoff"]["id"] = "C/D"
    first_source = _write_json(tmp_path / "first.json", first_payload)
    first = generate_signoff(first_source, tmp_path / "out")
    original = first.docx_path.read_bytes()

    second_payload = _fixture_payload("huntington_hospital_20260805.json")
    second_payload["site"]["code"] = "A_B"
    second_payload["signoff"]["id"] = "C_D"
    second_source = _write_json(tmp_path / "second.json", second_payload)
    with pytest.raises(SignoffValidationError, match="safe-slug collision"):
        generate_signoff(second_source, tmp_path / "out")
    assert first.docx_path.read_bytes() == original


def test_existing_lock_fails_closed(tmp_path: Path) -> None:
    payload = _fixture_payload("huntington_hospital_20260805.json")
    source = _write_json(tmp_path / "locked.json", payload)
    site_dir = tmp_path / "out" / "HUNTINGTON-HOSPITAL"
    site_dir.mkdir(parents=True)
    lock = site_dir / ".HUNTINGTON-STOCK-20260805-001.lock"
    lock.write_text("owned", encoding="utf-8")
    with pytest.raises(SignoffValidationError, match="another generation run"):
        generate_signoff(source, tmp_path / "out")
    assert lock.read_text(encoding="utf-8") == "owned"


def test_cli_runs_without_pythonpath_in_canonical_tree() -> None:
    env = os.environ.copy()
    env.pop("PYTHONPATH", None)
    output_root = ROOT / "Outputs" / "delivery-signoff" / f"pytest-{uuid4().hex}"
    try:
        completed = subprocess.run(
            [
                sys.executable,
                "scripts/generate_delivery_signoff.py",
                str(FIXTURES / "huntington_hospital_20260805.json"),
                "--output-root",
                str(output_root),
            ],
            cwd=ROOT,
            env=env,
            capture_output=True,
            text=True,
            check=False,
        )
        assert completed.returncode == 0, completed.stderr
        assert "PASS: delivery sign-off generated" in completed.stdout
        assert list(output_root.rglob("delivery-signoff-artifact-manifest.json"))
    finally:
        shutil.rmtree(output_root, ignore_errors=True)


def test_cli_rejects_noncanonical_output(tmp_path: Path) -> None:
    completed = subprocess.run(
        [
            sys.executable,
            "scripts/generate_delivery_signoff.py",
            str(FIXTURES / "huntington_hospital_20260805.json"),
            "--output-root",
            str(tmp_path / "not-canonical"),
        ],
        cwd=ROOT,
        capture_output=True,
        text=True,
        check=False,
    )
    assert completed.returncode == 1
    assert completed.stderr.startswith("FAIL:")
    assert not (tmp_path / "not-canonical").exists()


def _serial_payload(serials: list[str]) -> dict:
    payload = _fixture_payload("huntington_hospital_20260805.json")
    payload["equipment_rows"] = [
        {"equipment_type": "Neuron", "model_or_part": "", "color_or_variant": "", "quantity": len(serials)}
    ]
    payload["serialized_assets"] = [
        {
            "asset_type": "Neuron",
            "identifiers": [
                {"serial_number": serial, "mac_address": ""}
                for serial in serials
            ],
        }
    ]
    return payload


def _fixture_payload(name: str) -> dict:
    return json.loads((FIXTURES / name).read_text(encoding="utf-8"))


def _write_json(path: Path, payload: dict) -> Path:
    path.write_text(json.dumps(payload), encoding="utf-8")
    return path


def _manifest(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _docx_xml(path: Path) -> str:
    import zipfile

    with zipfile.ZipFile(path) as archive:
        return archive.read("word/document.xml").decode("utf-8")
