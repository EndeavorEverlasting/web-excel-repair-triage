"""
tests/test_invoice_parser.py
----------------------------
Unit and integration tests for triage.invoice_parser.

Coverage
--------
  T1  _classify_line — trucking keywords (AAA Disposal style)
  T2  _classify_line — courier keywords (NYM Courier style)
  T3  _classify_line — labor keywords
  T4  _classify_line — prefix overrides body keyword (Labor: logistics team)
  T5  _classify_line — unknown description returns 'other'
  T6  _detect_vendor — AAA Disposal recognised in header area
  T7  _detect_vendor — NYM Courier / New York Minute recognised
  T8  _detect_vendor — unknown vendor returns 'Unknown Vendor'
  T9  _parse_line_items — tab-separated row parsed correctly
  T10 _parse_line_items — sequential 5-line group parsed correctly
  T11 _parse_line_items — summary-only lines do NOT produce line items
  T12 parse_invoice — missing file raises InvoiceParseError
  T13 parse_invoice — non-.docx extension raises InvoiceParseError
  T14 parse_invoice — AAA Disposal format: vendor, category, total, line items
  T15 parse_invoice — NYM Courier format: vendor, category, total, line items
  T16 parse_invoice — unrecognised structure (total but no line items) raises error
  T17 parse_invoice — empty file raises InvoiceParseError
  T18 _classify_invoice — trucking invoice picked when trucking dominates by amount
  T19 parse_invoice — currency extracted correctly
  T20 parse_invoice — PO number extracted correctly
"""
from __future__ import annotations

import os
import pytest
from pathlib import Path

from triage.invoice_parser import (
    InvoiceParseError,
    _classify_line,
    _classify_invoice,
    _detect_vendor,
    _parse_line_items,
    parse_invoice,
)


# ── _classify_line ────────────────────────────────────────────────────────────

class TestClassifyLine:
    def test_trucking_keyword(self):
        assert _classify_line("Truck delivery — 2 loads") == "trucking"

    def test_disposal_keyword(self):
        assert _classify_line("Waste disposal service") == "trucking"

    def test_courier_keyword(self):
        assert _classify_line("Courier pickup — rush") == "courier"

    def test_nym_keyword(self):
        assert _classify_line("NYM Courier run — downtown") == "courier"

    def test_labor_keyword(self):
        assert _classify_line("Labor: 4 crew members on site") == "labor"

    def test_labor_prefix_beats_logistics(self):
        """'Labor: ...' prefix wins even when 'logistics' appears in the body."""
        assert _classify_line("Labor: 3-person logistics team") == "labor"

    def test_tech_prefix(self):
        assert _classify_line("Technician installation — full day") == "labor"

    def test_unknown_returns_other(self):
        assert _classify_line("Office supplies") == "other"

    def test_empty_string(self):
        assert _classify_line("") == "other"


# ── _detect_vendor ────────────────────────────────────────────────────────────

class TestDetectVendor:
    def test_aaa_disposal(self):
        text = "AAA Disposal\nInvoice No: INV-001\nService Date: Apr 01, 2026"
        assert _detect_vendor(text) == "AAA Disposal"

    def test_nym_courier_full_name(self):
        text = "New York Minute Courier\nInvoice #: NYM-2026-04"
        assert _detect_vendor(text) == "NYM Courier"

    def test_nym_courier_abbrev(self):
        text = "NYM Courier\nPO: 123456"
        assert _detect_vendor(text) == "NYM Courier"

    def test_unknown_vendor(self):
        text = "Generic Supplies Inc\nInvoice: 999"
        assert _detect_vendor(text) == "Unknown Vendor"

    def test_body_text_does_not_hijack(self):
        """'Cybernet deployment' in body should not override vendor from header."""
        text = "AAA Disposal\nService for Cybernet deployment"
        assert _detect_vendor(text) == "AAA Disposal"


# ── _parse_line_items ─────────────────────────────────────────────────────────

class TestParseLineItems:
    def test_tab_separated_row(self):
        text = "Billing Summary\nItem Description\tQty\tUnit\tRate\tAmount\nTruck rental\t1\tday\t500.00\t500.00\nSubtotal"
        items = _parse_line_items(text)
        assert len(items) >= 1
        item = items[0]
        assert "truck" in item["description"].lower() or item["amount"] == 500.0
        assert item["amount"] == 500.0

    def test_sequential_five_line_group(self):
        lines = [
            "Billing Summary",
            "Item Description",
            "Waste Hauling",
            "2",
            "trips",
            "350.00",
            "700.00",
            "Subtotal",
        ]
        text = "\n".join(lines)
        items = _parse_line_items(text)
        assert len(items) >= 1
        item = items[0]
        assert item["description"] == "Waste Hauling"
        assert item["amount"] == pytest.approx(700.0)
        assert item["qty"] == pytest.approx(2.0)
        assert item["unit"] == "trips"

    def test_summary_lines_not_parsed_as_items(self):
        """Lines that are invoice totals (Total, Subtotal) must not become line items."""
        text = "Subtotal\nTotal Due 1000.00\n"
        items = _parse_line_items(text)
        amounts = [i["amount"] for i in items]
        assert 1000.0 not in amounts, "Summary line was wrongly parsed as a line item"


# ── parse_invoice — error cases ───────────────────────────────────────────────

class TestParseInvoiceErrors:
    def test_missing_file_raises(self, tmp_path):
        with pytest.raises(InvoiceParseError, match="not found"):
            parse_invoice(str(tmp_path / "nonexistent.docx"))

    def test_wrong_extension_raises(self, tmp_path):
        p = tmp_path / "invoice.pdf"
        p.write_text("data")
        with pytest.raises(InvoiceParseError, match=".docx"):
            parse_invoice(str(p))

    def test_unrecognised_structure_raises(self, tmp_path):
        """A .docx with a total but no parseable line items must raise InvoiceParseError."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Some Vendor")
        doc.add_paragraph("Total 9999.00")
        path = str(tmp_path / "bad_invoice.docx")
        doc.save(path)
        with pytest.raises(InvoiceParseError):
            parse_invoice(path)

    def test_empty_document_raises(self, tmp_path):
        from docx import Document
        doc = Document()
        path = str(tmp_path / "empty.docx")
        doc.save(path)
        with pytest.raises(InvoiceParseError, match="empty|no readable text|amount"):
            parse_invoice(path)


# ── parse_invoice — AAA Disposal format ──────────────────────────────────────

class TestParseInvoiceAAA:
    @pytest.fixture(scope="class")
    def aaa_invoice(self, tmp_path_factory):
        from docx import Document
        tmp = tmp_path_factory.mktemp("aaa")
        doc = Document()
        doc.add_paragraph("AAA Disposal")
        doc.add_paragraph("Invoice No: INV-AAA-001")
        doc.add_paragraph("PO No: 176759")
        doc.add_paragraph("Service Date: April 01, 2026")
        doc.add_paragraph("Prepared for: Agilant Solutions")
        doc.add_paragraph("Prepared by: AAA Disposal Billing")
        doc.add_paragraph("USD")
        doc.add_paragraph("Billing Summary")
        doc.add_paragraph("Item Description")
        doc.add_paragraph("Waste Hauling")
        doc.add_paragraph("2")
        doc.add_paragraph("trips")
        doc.add_paragraph("350.00")
        doc.add_paragraph("700.00")
        doc.add_paragraph("Subtotal")
        doc.add_paragraph("Subtotal 700.00")
        doc.add_paragraph("Total Due 700.00")
        path = str(tmp / "aaa_invoice.docx")
        doc.save(path)
        return parse_invoice(path)

    def test_vendor(self, aaa_invoice):
        assert aaa_invoice["vendor"] == "AAA Disposal"

    def test_cost_category_trucking(self, aaa_invoice):
        assert aaa_invoice["cost_category"] == "trucking"

    def test_invoice_number(self, aaa_invoice):
        assert aaa_invoice["invoice_number"] == "INV-AAA-001"

    def test_po_number(self, aaa_invoice):
        assert aaa_invoice["po_number"] == "176759"

    def test_total_positive(self, aaa_invoice):
        assert aaa_invoice["total"] is not None
        assert aaa_invoice["total"] > 0

    def test_line_items_present(self, aaa_invoice):
        assert len(aaa_invoice["line_items"]) >= 1

    def test_required_keys_present(self, aaa_invoice):
        required = {
            "vendor", "invoice_number", "po_number", "service_date",
            "line_items", "subtotal", "total", "cost_category", "currency",
        }
        for key in required:
            assert key in aaa_invoice, f"Missing key: {key}"


# ── parse_invoice — NYM Courier format ───────────────────────────────────────

class TestParseInvoiceNYM:
    @pytest.fixture(scope="class")
    def nym_invoice(self, tmp_path_factory):
        from docx import Document
        tmp = tmp_path_factory.mktemp("nym")
        doc = Document()
        doc.add_paragraph("NYM Courier")
        doc.add_paragraph("Invoice No: NYM-2026-042")
        doc.add_paragraph("PO No: 176760")
        doc.add_paragraph("Service Date: April 05, 2026")
        doc.add_paragraph("Prepared for: Agilant Solutions")
        doc.add_paragraph("Prepared by: NYM Billing Dept")
        doc.add_paragraph("USD")
        doc.add_paragraph("Billing Summary")
        doc.add_paragraph("Item Description")
        doc.add_paragraph("Courier delivery — express")
        doc.add_paragraph("3")
        doc.add_paragraph("trips")
        doc.add_paragraph("120.00")
        doc.add_paragraph("360.00")
        doc.add_paragraph("Subtotal")
        doc.add_paragraph("Subtotal 360.00")
        doc.add_paragraph("Total Due 360.00")
        path = str(tmp / "nym_invoice.docx")
        doc.save(path)
        return parse_invoice(path)

    def test_vendor(self, nym_invoice):
        assert nym_invoice["vendor"] == "NYM Courier"

    def test_cost_category_courier(self, nym_invoice):
        assert nym_invoice["cost_category"] == "courier"

    def test_invoice_number(self, nym_invoice):
        assert nym_invoice["invoice_number"] == "NYM-2026-042"

    def test_total_positive(self, nym_invoice):
        assert nym_invoice["total"] is not None
        assert nym_invoice["total"] > 0

    def test_line_items_present(self, nym_invoice):
        assert len(nym_invoice["line_items"]) >= 1

    def test_currency(self, nym_invoice):
        assert nym_invoice["currency"] == "USD"

    def test_service_date(self, nym_invoice):
        assert nym_invoice["service_date"] is not None
        assert "April" in nym_invoice["service_date"] or "2026" in nym_invoice["service_date"]


# ── _classify_invoice ─────────────────────────────────────────────────────────

class TestClassifyInvoice:
    def test_trucking_dominates(self):
        items = [
            {"description": "Truck delivery", "category": "trucking", "amount": 800.0},
            {"description": "Labor: one tech", "category": "labor",    "amount": 200.0},
        ]
        assert _classify_invoice(items) == "trucking"

    def test_courier_dominates(self):
        items = [
            {"description": "Courier run", "category": "courier", "amount": 500.0},
            {"description": "Other charge", "category": "other",   "amount": 50.0},
        ]
        assert _classify_invoice(items) == "courier"

    def test_empty_items_with_vendor_hint(self):
        assert _classify_invoice([], vendor_hint="nym courier run") == "courier"
