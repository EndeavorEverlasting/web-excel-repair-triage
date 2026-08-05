"""DOCX construction for delivery sign-offs."""
from __future__ import annotations

import math
import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .schema import MINIMUM_FONT_POINTS


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any, *, top: int = 40, start: int = 60, bottom: int = 40, end: int = 60) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_keep_with_next(paragraph: Any) -> None:
    paragraph._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))


def _format_run(run: Any, *, size: float = MINIMUM_FONT_POINTS, bold: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold


def _write_cell(cell: Any, text: str, *, bold: bool = False, size: float = MINIMUM_FONT_POINTS, align: Any = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    _format_run(paragraph.add_run(text), size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)


def _add_info_table(document: Document, spec: dict[str, Any]) -> None:
    table = document.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.75, 3.15, 0.95, 2.35]
    values = [
        ("Site", spec["site"]["name"], "Site Code", spec["site"]["code"]),
        ("Sign-Off ID", spec["signoff"]["id"], "Delivery Date", spec["signoff"]["delivery_date"] or "________________"),
    ]
    for row, row_values in zip(table.rows, values, strict=True):
        for index, (cell, width, value) in enumerate(zip(row.cells, widths, row_values, strict=True)):
            cell.width = Inches(width)
            _write_cell(cell, value, bold=index % 2 == 0)
            if index % 2 == 0:
                _set_cell_shading(cell, "D9E2F3")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_equipment_table(document: Document, rows: list[dict[str, Any]]) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(2)
    heading.paragraph_format.space_after = Pt(2)
    _set_keep_with_next(heading)
    _format_run(heading.add_run("Equipment / Stock Receipt"), size=10.5, bold=True)

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    widths = [2.45, 1.65, 1.35, 0.65, 0.95]
    headers = ["Item Description", "Model / Part", "Color / Variant", "Qty", "Initials"]
    for cell, width, header in zip(table.rows[0].cells, widths, headers, strict=True):
        cell.width = Inches(width)
        _write_cell(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(cell, "D9E2F3")
    _set_repeat_table_header(table.rows[0])
    for item in rows:
        row = table.add_row()
        values = [item["equipment_type"], item["model_or_part"], item["color_or_variant"], str(item["quantity"]), ""]
        for index, (cell, width, value) in enumerate(zip(row.cells, widths, values, strict=True)):
            cell.width = Inches(width)
            _write_cell(cell, value, align=WD_ALIGN_PARAGRAPH.CENTER if index in (3, 4) else None)


def _format_identifier(item: dict[str, str]) -> str:
    return f"{item['serial_number']} / {item['mac_address']}" if item["mac_address"] else item["serial_number"]


def _add_serial_group(document: Document, group: dict[str, Any], columns: int = 3) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(3)
    heading.paragraph_format.space_after = Pt(1)
    _set_keep_with_next(heading)
    _format_run(heading.add_run(f"{group['asset_type']} Serial Verification ({len(group['identifiers'])})"), size=10, bold=True)

    identifiers = group["identifiers"]
    rows_needed = math.ceil(len(identifiers) / columns)
    table = document.add_table(rows=1 + rows_needed, cols=columns * 3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for column in range(columns):
        cells = table.rows[0].cells[column * 3 : column * 3 + 3]
        for cell, text, width in zip(cells, ("#", "Serial / MAC", "Mark"), (0.28, 1.75, 0.42), strict=True):
            cell.width = Inches(width)
            _write_cell(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_shading(cell, "D9E2F3")
    _set_repeat_table_header(table.rows[0])

    for row_index in range(rows_needed):
        row = table.rows[row_index + 1]
        for column in range(columns):
            item_index = row_index + column * rows_needed
            cells = row.cells[column * 3 : column * 3 + 3]
            values = (
                [str(item_index + 1), _format_identifier(identifiers[item_index]), "[  ]"]
                if item_index < len(identifiers)
                else ["", "", ""]
            )
            for item_column, (cell, text, width) in enumerate(zip(cells, values, (0.28, 1.75, 0.42), strict=True)):
                cell.width = Inches(width)
                _write_cell(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER if item_column != 1 else None)


def _add_annotation_and_acceptance(document: Document, spec: dict[str, Any], *, notes_height: float) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(3)
    heading.paragraph_format.space_after = Pt(1)
    _set_keep_with_next(heading)
    _format_run(heading.add_run("Exceptions / Field Notes"), size=9.5, bold=True)

    notes = document.add_table(rows=1, cols=1)
    notes.style = "Table Grid"
    notes.rows[0].height = Inches(notes_height)
    notes.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    _write_cell(notes.cell(0, 0), "")

    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(3)
    heading.paragraph_format.space_after = Pt(1)
    _set_keep_with_next(heading)
    _format_run(heading.add_run("Receipt Acceptance"), size=9.5, bold=True)

    recipient = spec["recipient"]
    table = document.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    values = [
        ("POC Name", recipient["name"] or "________________________", "Title", recipient["title"] or "________________"),
        ("Building / Room", recipient["building_room"] or "________________________", "Phone", recipient["phone"] or "________________"),
        ("Signature", "", "Date", ""),
    ]
    widths = [1.05, 2.65, 0.75, 2.55]
    for row, row_values in zip(table.rows, values, strict=True):
        for index, (cell, value, width) in enumerate(zip(row.cells, row_values, widths, strict=True)):
            cell.width = Inches(width)
            _write_cell(cell, value, bold=index % 2 == 0)
            if index % 2 == 0:
                _set_cell_shading(cell, "D9E2F3")


def build_document(spec: dict[str, Any], output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    serial_total = sum(len(group["identifiers"]) for group in spec["serialized_assets"])
    if serial_total > 0 or len(spec["equipment_rows"]) > 18:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.35)
    section.header_distance = section.footer_distance = Inches(0.15)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(MINIMUM_FONT_POINTS)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    _format_run(title.add_run(spec["signoff"]["title"]), size=14, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(3)
    _format_run(subtitle.add_run(spec["signoff"]["subtitle"] or spec["site"]["name"]), size=10.5, bold=True)

    _add_info_table(document, spec)
    _add_equipment_table(document, spec["equipment_rows"])
    for group in spec["serialized_assets"]:
        _add_serial_group(document, group)
    notes_height = 0.55 if spec["serialized_assets"] else max(1.0, min(5.8, 6.8 - (0.23 * len(spec["equipment_rows"]))))
    _add_annotation_and_acceptance(document, spec, notes_height=notes_height)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _format_run(footer.add_run(f"{spec['signoff']['id']} | Editable, unprotected sign-off"), size=7.5)
    document.core_properties.title = spec["signoff"]["title"]
    document.core_properties.subject = spec["site"]["name"]
    document.core_properties.comments = "Generated by Triage delivery sign-off generator"
    document.save(output_path)


def docx_text(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
    return re.sub(r"<[^>]+>", " ", xml)
